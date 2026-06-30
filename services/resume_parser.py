"""
简历解析服务（增强版）

参照参考项目的多级简历处理管线，从基础文本提取升级为完整的：
  ┌─ 第1层：文件分类 → 直接文本 / OCR
  ├─ 第2a层：DOCX → python-docx + zipfile/XML 兜底 + 表格提取
  ├─ 第2b层：TXT/MD → 多编码探测
  ├─ 第2c层：PDF → PyMuPDF(优先) / PyPDF2(兜底) + 扫描版转OCR
  ├─ 第2d层：图片 → EasyOCR 本地识别
  ├─ 第3层：文本清洗 → 零宽字符/BOM/控制字符/空行合并
  └─ 第4层：智能摘要 → 段落边界截断 + 预览

支持格式: PDF(文本型+扫描版), DOCX, JPG/PNG/BMP/WEBP, TXT, MD
"""
import os
import re
from pathlib import Path
from typing import Optional, Dict, Tuple

# ── 常量 ──

# 需要 OCR 的文件扩展名（PDF 也走 OCR 优先，因为可能是扫描版）
OCR_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".webp"}

# 可直接读文本的扩展名
DIRECT_TEXT_EXTENSIONS = {".docx", ".doc", ".md", ".markdown", ".txt"}

# 支持的所有格式
SUPPORTED_EXTENSIONS = OCR_EXTENSIONS | DIRECT_TEXT_EXTENSIONS

# 需要清洗的不可见字符
# 注意: 零宽字符在源码中不可见，用 Unicode 转义保证可移植性
INVISIBLE_CHARS = [
    ("﻿", ""),      # BOM 头
    ("​", ""),      # 零宽空格
    ("‌", ""),      # 零宽非连接符
    ("‍", ""),      # 零宽连接符
    ("\xa0", " "),        # 不间断空格 → 普通空格
    ("\r\n", "\n"),       # Windows 换行
    ("\r", "\n"),          # Mac 老式换行
]

# OCR 错误标记
OCR_ERROR_MARKERS = (
    "错误:", "ocr api", "ocr 调用异常",
    "api 令牌未配置", "无法连接到 ocr api",
)


# ═══════════════════════════════════════════════
# 公开 API
# ═══════════════════════════════════════════════

def extract_resume_text(file_path: str, use_ocr: bool = True) -> str:
    """
    从简历文件中提取文本内容（增强版入口）

    Args:
        file_path: 简历文件路径
        use_ocr: 是否启用 OCR（默认开启）

    Returns:
        清洗后的文本内容
    """
    if not os.path.exists(file_path):
        return ""

    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return f"[错误] 不支持的格式: {ext}。支持: PDF, DOCX, TXT, MD, JPG, PNG, BMP, WEBP"

    try:
        # 第1层：文件分类 + 第2层：文本提取
        if ext in DIRECT_TEXT_EXTENSIONS:
            raw_text = _extract_direct(file_path, ext)
        elif use_ocr:
            raw_text = _extract_via_ocr(file_path)
        else:
            return "[错误] OCR 未启用，当前文件类型需要 OCR 才能解析"

        if not raw_text or not raw_text.strip():
            return "[错误] 未能从文件中提取到任何文本"

        # 检查 OCR 错误
        lowered = raw_text.lower()
        if any(m in lowered for m in OCR_ERROR_MARKERS):
            return raw_text  # 保留原始错误信息

        # 第3层：文本清洗
        clean = sanitize_text(raw_text, ext=ext)

        if not clean or not clean.strip():
            return "[错误] 清洗后无可用的文本内容"

        return clean

    except Exception as e:
        return f"[错误] 解析失败 ({ext}): {e}"


def extract_resume_full(file_path: str) -> Dict:
    """
    完整提取简历（返回结构化信息）

    Returns:
        {
            "success": bool,
            "text": str,
            "preview": str,
            "char_count": int,
            "source": str,  # "docx" / "pdf_text" / "pdf_ocr" / "image_ocr" / "text"
        }
    """
    text = extract_resume_text(file_path)
    ext = Path(file_path).suffix.lower()

    success = not text.startswith("[错误]") and len(text.strip()) >= 20

    return {
        "success": success,
        "text": text if success else "",
        "preview": get_resume_preview(text) if success else text,
        "char_count": len(text) if success else 0,
        "source": _source_label(ext),
        "error": "" if success else text,
    }


def sanitize_text(text: str, ext: str = "") -> str:
    """
    文本清洗（参照参考项目 _sanitize_resume_text）

    处理流程:
    1. 替换不可见字符（BOM/零宽/控制字符）
    2. 统一换行符
    3. Markdown 格式符号剥离（仅对 .md/.markdown）
    4. 合并连续空白行
    5. 去除水平标尺线
    """
    if not text:
        return ""

    # Step 1: 不可见字符替换
    normalized = text
    for old, new in INVISIBLE_CHARS:
        normalized = normalized.replace(old, new)

    # Step 2: Markdown 特殊处理
    if ext in (".md", ".markdown"):
        # 去 HTML 注释
        normalized = re.sub(r"<!--.*?-->", "", normalized, flags=re.S)
        # 去图片引用
        normalized = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", normalized)
        # 去链接保留文字
        normalized = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", normalized)

    # Step 3: 逐行清洗
    lines = []
    prev_blank = False

    for raw_line in normalized.split("\n"):
        # Tab 转空格
        line = raw_line.replace("\t", "    ").strip()

        # Markdown 格式符号剥离
        if ext in (".md", ".markdown"):
            line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)   # 标题
            line = re.sub(r"^\s*>+\s*", "", line)             # 引用
            line = re.sub(r"`([^`]*)`", r"\1", line)          # 行内代码
            line = line.replace("**", "").replace("__", "")   # 加粗

        # 水平标尺线跳过
        if re.fullmatch(r"[-=_*~]{3,}", line):
            continue

        # 多余的空白符 → 单个空格
        line = re.sub(r"\s+", " ", line).strip()

        if not line:
            if prev_blank:
                continue
            lines.append("")
            prev_blank = True
            continue

        lines.append(line)
        prev_blank = False

    return "\n".join(lines).strip()


def summarize_resume(text: str, max_chars: int = 2000) -> str:
    """
    简历文本摘要（在段落边界处智能截断）

    避免在句子中间截断，优先在段落末尾、句子末尾处断开。
    """
    text = text.strip()
    if len(text) <= max_chars:
        return text

    truncated = text[:max_chars]
    # 按优先级找断点：段落 > 句号 > 换行
    for delimiter in ("\n\n", "\n", "。", "；"):
        last_break = truncated.rfind(delimiter)
        if last_break > max_chars * 0.5:
            truncated = truncated[:last_break]
            break

    return truncated + "\n\n[简历内容已截断，完整内容已保存]"


def get_resume_preview(text: str, max_chars: int = 300) -> str:
    """获取简历预览（用于列表展示）"""
    text = text.strip()
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    preview = " | ".join(lines[:3])
    if len(preview) > max_chars:
        preview = preview[:max_chars] + "..."
    return preview or "(无内容)"


# ═══════════════════════════════════════════════
# 第2a层：DOCX 提取（python-docx + zipfile/XML 兜底）
# ═══════════════════════════════════════════════

def _extract_docx(path: Path) -> str:
    """DOCX 文本提取：python-docx 优先，zipfile/XML 兜底"""
    # 主方案：python-docx
    try:
        from docx import Document
        doc = Document(str(path))
        blocks = []

        # 段落
        for para in doc.paragraphs:
            t = _normalize_text(para.text)
            if t:
                blocks.append(t)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                row_parts = [
                    _normalize_text(cell.text)
                    for cell in row.cells
                    if _normalize_text(cell.text)
                ]
                if row_parts:
                    blocks.append(" | ".join(row_parts))

        result = "\n\n".join(blocks).strip()
        if result:
            return result

    except Exception:
        pass

    # 兜底方案：zipfile + XML 直接解析
    try:
        import zipfile
        from xml.etree import ElementTree as ET

        WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        with zipfile.ZipFile(str(path)) as archive:
            if "word/document.xml" not in archive.namelist():
                return ""

            doc_xml = archive.read("word/document.xml")
            root = ET.fromstring(doc_xml)

            blocks = []
            body = root.find(f".//{{{WORD_NS}}}body")
            nodes = list(body) if body is not None else list(root.iter())

            for node in nodes:
                tag = node.tag.split("}", 1)[-1] if "}" in node.tag else node.tag

                if tag == "p":
                    texts = []
                    for t_elem in node.iter():
                        t_tag = t_elem.tag.split("}", 1)[-1] if "}" in t_elem.tag else t_elem.tag
                        if t_tag == "t" and t_elem.text:
                            texts.append(t_elem.text)
                    line = "".join(texts).strip()
                    if line:
                        blocks.append(line)

                elif tag == "tbl":
                    for row in node.findall(f".//{{{WORD_NS}}}tr"):
                        cells = []
                        for cell in row.findall(f".//{{{WORD_NS}}}tc"):
                            cell_texts = []
                            for p in cell.findall(f".//{{{WORD_NS}}}p"):
                                p_texts = []
                                for t_elem in p.iter():
                                    t_tag = t_elem.tag.split("}", 1)[-1] if "}" in t_elem.tag else t_elem.tag
                                    if t_tag == "t" and t_elem.text:
                                        p_texts.append(t_elem.text)
                                if p_texts:
                                    cell_texts.append("".join(p_texts))
                            if cell_texts:
                                cells.append(" / ".join(cell_texts))
                        if cells:
                            blocks.append(" | ".join(cells))

            # 去重
            seen = set()
            deduped = []
            for b in blocks:
                if b not in seen:
                    deduped.append(b)
                    seen.add(b)

            return "\n\n".join(deduped).strip()

    except Exception:
        return ""


# ═══════════════════════════════════════════════
# 第2b层：TXT/MD 多编码提取
# ═══════════════════════════════════════════════

def _extract_text_file(path: Path) -> str:
    """文本文件读取：多编码依次尝试"""
    encodings = ("utf-8-sig", "utf-8", "gb18030", "gbk")

    for enc in encodings:
        try:
            return _normalize_text(path.read_text(encoding=enc))
        except (UnicodeDecodeError, Exception):
            continue

    # 最终兜底：忽略无法解码的字符
    try:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    except Exception:
        return ""


# ═══════════════════════════════════════════════
# 第2c层：PDF 提取（PyMuPDF 优先 + OCR 扫描版兜底）
# ═══════════════════════════════════════════════

def _extract_pdf(path: Path) -> str:
    """PDF 提取：PyMuPDF 优先，PyPDF2 兜底，扫描版转 OCR"""
    text = ""

    # 方案1：PyMuPDF（提取质量最好）
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                pages.append(page_text.strip())
        doc.close()
        text = "\n\n".join(pages)
    except ImportError:
        pass
    except Exception:
        pass

    # 方案2：PyPDF2（兜底）
    if not text:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages.append(page_text.strip())
            text = "\n\n".join(pages)
        except Exception:
            pass

    # 方案3：如果文本提取结果过少（可能是扫描版 PDF），用 OCR
    if len(text.strip()) < 50:
        ocr_text = _ocr_image(str(path))
        if ocr_text and len(ocr_text) > len(text):
            text = ocr_text

    return text


# ═══════════════════════════════════════════════
# 第2d层：OCR 识别（EasyOCR 本地引擎）
# ═══════════════════════════════════════════════

_easyocr_reader = None  # 惰性加载，避免启动时消耗内存


def _get_ocr_reader():
    """惰性加载 EasyOCR reader（单例，约 100MB 模型首次下载）"""
    global _easyocr_reader
    if _easyocr_reader is None:
        try:
            import easyocr
            # 中文 + 英文，使用 GPU 如果可用
            _easyocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=True, verbose=False)
        except ImportError:
            return None
    return _easyocr_reader


def _ocr_image(file_path: str) -> str:
    """对图片/扫描PDF执行 OCR"""
    reader = _get_ocr_reader()
    if reader is None:
        return ""

    try:
        results = reader.readtext(str(file_path), detail=0)
        return "\n".join(results)
    except Exception:
        return ""


# ═══════════════════════════════════════════════
# 第1层：文件分类 + 第2层调度
# ═══════════════════════════════════════════════

def _extract_direct(file_path: str, ext: str) -> str:
    """直接文本提取入口"""
    path = Path(file_path)

    if ext in (".docx", ".doc"):
        return _extract_docx(path)
    else:
        return _extract_text_file(path)


def _extract_via_ocr(file_path: str) -> str:
    """OCR 提取入口"""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(Path(file_path))
    else:
        return _ocr_image(file_path)


def _source_label(ext: str) -> str:
    """返回来源标签"""
    labels = {
        ".docx": "Word文档", ".doc": "Word文档",
        ".md": "Markdown", ".markdown": "Markdown",
        ".txt": "文本文件",
        ".pdf": "PDF(OCR)", ".jpg": "图片OCR", ".jpeg": "图片OCR",
        ".png": "图片OCR", ".bmp": "图片OCR", ".webp": "图片OCR",
    }
    return labels.get(ext, "未知来源")


def _normalize_text(value: object) -> str:
    """标准化文本：统一换行，去除行尾空格"""
    text = str(value or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    return "\n".join(lines).strip()
