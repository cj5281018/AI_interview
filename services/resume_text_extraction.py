"""
简历文本提取服务 — 完整复刻参考项目

处理管线：
  ┌─ 文件分类：OCR_EXTENSIONS vs DIRECT_TEXT_EXTENSIONS
  ├─ DOCX：python-docx(主) → zipfile+XML 解析(兜底) → 页眉/页脚/脚注
  ├─ TXT/MD：多编码依次探测（utf-8-sig → utf-8 → gb18030 → gbk）
  ├─ 文本清洗：7种不可见字符替换 + Markdown 格式剥离 + 空行合并
  └─ 错误检测：OCR 错误标记检测 + 空文本检测
"""
from __future__ import annotations

from pathlib import Path
import re
from typing import Callable
from xml.etree import ElementTree as ET
import zipfile

try:
    from docx import Document
except ImportError:
    Document = None

# ── 常量 ──
OCR_RESUME_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".webp", ".heic", ".heif"}
DIRECT_TEXT_EXTENSIONS = {".docx", ".md", ".markdown", ".txt"}
SUPPORTED_RESUME_EXTENSIONS = OCR_RESUME_EXTENSIONS | DIRECT_TEXT_EXTENSIONS

INVALID_TEXT_MARKERS = (
    "错误:", "ocr api", "ocr 调用异常",
    "api 令牌未配置", "无法连接到 ocr api", "无法连接到ocr",
)
EMPTY_TEXT_MARKERS = (
    "未提取到任何有效文本", "未能从图片中提取到任何有效文本",
)
DIRECT_SOURCE_LABELS = {
    ".docx": "Word 文档", ".md": "Markdown 文档",
    ".markdown": "Markdown 文档", ".txt": "文本文件",
}
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DOCX_SUPPLEMENTAL_PART_PATTERN = re.compile(
    r"^word/(?:header\d*|footer\d*|footnotes|endnotes|comments\d*)\.xml$"
)


class ResumeExtractionError(Exception):
    """简历提取基础异常"""

class UnsupportedResumeFormatError(ResumeExtractionError):
    """不支持的简历格式"""

class ResumeOcrUnavailableError(ResumeExtractionError):
    """OCR 不可用"""


# ═══════════════════════════════════════
# 公开 API
# ═══════════════════════════════════════

def get_resume_extension(file_name_or_path: str) -> str:
    return Path(str(file_name_or_path or "")).suffix.lower()


def ensure_supported_resume_extension(file_name_or_path: str) -> str:
    ext = get_resume_extension(file_name_or_path)
    if ext in SUPPORTED_RESUME_EXTENSIONS:
        return ext
    if ext == ".doc":
        raise UnsupportedResumeFormatError("暂不支持旧版 .doc，请先另存为 .docx 后再上传。")
    if not ext:
        raise UnsupportedResumeFormatError("无法识别简历文件类型，请上传带扩展名的文件。")
    raise UnsupportedResumeFormatError(
        f"不支持的简历格式: {ext}。当前支持 PDF、图片、DOCX、MD、TXT。"
    )


def resume_requires_ocr(file_name_or_path: str) -> bool:
    return ensure_supported_resume_extension(file_name_or_path) in OCR_RESUME_EXTENSIONS


def extract_resume_content(
    file_path: str,
    *,
    include_images: bool = False,
    ocr_available: bool = True,
    ocr_text_loader: Callable[..., str] | None = None,
    ocr_full_loader: Callable[..., dict] | None = None,
    use_preprocessing: bool = True,
    is_screen_capture: bool = False,
) -> dict:
    """主入口：与参考项目完全一致的签名和返回格式"""
    ext = ensure_supported_resume_extension(file_path)

    # ── 直接文本提取（DOCX/TXT/MD）──
    if ext in DIRECT_TEXT_EXTENSIONS:
        clean_text = _sanitize_resume_text(_extract_direct_text(file_path, ext), ext=ext)
        if not clean_text:
            return {
                "success": False, "mode": "direct_text",
                "source_label": DIRECT_SOURCE_LABELS.get(ext, "文本"),
                "text": "", "reusable_text": "", "raw_text": "", "images": {},
                "error_message": f"{DIRECT_SOURCE_LABELS.get(ext, '文本文件')}中未提取到有效文本。",
            }
        reusable_text = _format_direct_text_result(clean_text, ext)
        return {
            "success": True, "mode": "direct_text",
            "source_label": DIRECT_SOURCE_LABELS.get(ext, "文本"),
            "text": clean_text, "reusable_text": reusable_text,
            "raw_text": reusable_text, "images": {}, "error_message": "",
        }

    # ── OCR 提取（PDF/图片）──
    if not ocr_available:
        raise ResumeOcrUnavailableError("OCR 模块不可用，当前文件类型需要 OCR 才能解析。")

    if include_images:
        if ocr_full_loader is None:
            raise ResumeOcrUnavailableError("OCR 完整解析器未加载。")
        ocr_result = ocr_full_loader(
            image_path=file_path, use_preprocessing=use_preprocessing,
            is_screen_capture=is_screen_capture,
        ) or {}
        raw_text = str(ocr_result.get("text") or "").strip()
        images = ocr_result.get("images") if isinstance(ocr_result, dict) else {}
        if not isinstance(images, dict):
            images = {}
    else:
        if ocr_text_loader is None:
            raise ResumeOcrUnavailableError("OCR 解析器未加载。")
        raw_text = str(ocr_text_loader(
            image_path=file_path, use_preprocessing=use_preprocessing,
            is_screen_capture=is_screen_capture,
        ) or "").strip()
        images = {}

    clean_text = unwrap_resume_text(raw_text)
    success = _is_successful_extraction(raw_text, clean_text)
    return {
        "success": success, "mode": "ocr", "source_label": "OCR",
        "text": clean_text if success else "",
        "reusable_text": raw_text if success else "",
        "raw_text": raw_text,
        "images": images if success else {},
        "error_message": "" if success else (raw_text or "OCR 未提取到有效文本。"),
    }


# ═══════════════════════════════════════
# 直接文本提取
# ═══════════════════════════════════════

def _extract_direct_text(file_path: str, ext: str) -> str:
    path = Path(file_path)
    if ext == ".docx":
        return _extract_docx_text(path)
    return _read_text_file(path)


# ═══════════════════════════════════════
# DOCX 提取：python-docx 主方案 + zipfile/XML 兜底
# ═══════════════════════════════════════

def _extract_docx_text(path: Path) -> str:
    """DOCX 文本提取：python-docx 优先，zipfile/XML 兜底"""
    if Document is not None:
        try:
            document = Document(path)
            blocks: list[str] = []
            for paragraph in document.paragraphs:
                text = _normalize_text(paragraph.text)
                if text:
                    blocks.append(text)
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell_text for cell_text in
                        (_normalize_text(cell.text) for cell in row.cells)
                        if cell_text
                    )
                    if row_text:
                        blocks.append(row_text)
            extracted_text = "\n\n".join(blocks).strip()
            if extracted_text:
                return extracted_text
            fallback_text = _extract_docx_text_via_zip(path)
            if fallback_text:
                return fallback_text
            return ""
        except Exception:
            fallback_text = _extract_docx_text_via_zip(path)
            if fallback_text:
                return fallback_text
            raise ResumeExtractionError("读取 Word 文档失败")

    return _extract_docx_text_via_zip(path)


def _extract_docx_text_via_zip(path: Path) -> str:
    """zipfile + XML 直接解析 docx（终极兜底）"""
    try:
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml")
            root = ET.fromstring(document_xml)
            blocks = _extract_docx_blocks_from_root(root)
            if blocks:
                return "\n\n".join(blocks).strip()
            supplemental_blocks = _extract_docx_blocks_from_supplemental_parts(archive)
            return "\n\n".join(supplemental_blocks).strip()
    except KeyError:
        raise ResumeExtractionError("Word 文档缺少 word/document.xml，文件可能已损坏。")
    except zipfile.BadZipFile:
        raise ResumeExtractionError("Word 文档不是有效的 .docx 文件。")
    except ET.ParseError:
        raise ResumeExtractionError("解析 Word 文档 XML 失败。")
    except Exception as exc:
        raise ResumeExtractionError(f"读取 Word 文档失败: {exc}")


def _extract_docx_blocks_from_root(root: ET.Element) -> list[str]:
    body = root.find(f".//{{{WORD_NAMESPACE}}}body")
    if body is not None:
        return _extract_docx_blocks_from_container(body)
    paragraphs: list[str] = []
    for paragraph in root.iterfind(f".//{{{WORD_NAMESPACE}}}p"):
        paragraph_text = _collect_docx_paragraph_text(paragraph)
        if paragraph_text:
            paragraphs.append(paragraph_text)
    return _dedupe_docx_blocks(paragraphs)


def _extract_docx_blocks_from_container(container: ET.Element) -> list[str]:
    blocks: list[str] = []
    for child in list(container):
        tag = _strip_xml_namespace(child.tag)
        if tag == "p":
            paragraph_text = _collect_docx_paragraph_text(child)
            if paragraph_text:
                blocks.append(paragraph_text)
        elif tag == "tbl":
            blocks.extend(_collect_docx_table_rows(child))
    return _dedupe_docx_blocks(blocks)


def _extract_docx_blocks_from_supplemental_parts(archive: zipfile.ZipFile) -> list[str]:
    """提取页眉、页脚、脚注等补充部件"""
    blocks: list[str] = []
    for part_name in sorted(archive.namelist()):
        if not DOCX_SUPPLEMENTAL_PART_PATTERN.match(part_name):
            continue
        try:
            root = ET.fromstring(archive.read(part_name))
        except ET.ParseError:
            continue
        blocks.extend(_extract_docx_blocks_from_root(root))
    return _dedupe_docx_blocks(blocks)


def _dedupe_docx_blocks(blocks: list[str]) -> list[str]:
    deduped: list[str] = []
    seen: set[str] = set()
    for block in blocks:
        key = block.strip()
        if not key or key in seen:
            continue
        deduped.append(key)
        seen.add(key)
    return deduped


def _collect_docx_paragraph_text(node: ET.Element) -> str:
    fragments: list[str] = []
    for element in node.iter():
        tag = _strip_xml_namespace(element.tag)
        if tag == "t" and element.text:
            fragments.append(element.text)
        elif tag == "tab":
            fragments.append("\t")
        elif tag in {"br", "cr"}:
            fragments.append("\n")
    return _normalize_text("".join(fragments))


def _collect_docx_table_rows(table_node: ET.Element) -> list[str]:
    rows: list[str] = []
    for row in table_node.findall(f"./{{{WORD_NAMESPACE}}}tr"):
        cells: list[str] = []
        for cell in row.findall(f"./{{{WORD_NAMESPACE}}}tc"):
            cell_lines: list[str] = []
            for paragraph in cell.findall(f".//{{{WORD_NAMESPACE}}}p"):
                paragraph_text = _collect_docx_paragraph_text(paragraph)
                if paragraph_text:
                    cell_lines.append(paragraph_text)
            if cell_lines:
                cells.append(" / ".join(cell_lines))
        if cells:
            rows.append(" | ".join(cells))
    return rows


def _strip_xml_namespace(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


# ═══════════════════════════════════════
# TXT/MD：多编码文本读取
# ═══════════════════════════════════════

def _read_text_file(path: Path) -> str:
    encodings = ("utf-8-sig", "utf-8", "gb18030", "gbk")
    for encoding in encodings:
        try:
            return _normalize_text(path.read_text(encoding=encoding))
        except UnicodeDecodeError:
            continue
        except Exception:
            raise ResumeExtractionError(f"读取文本文件失败")
    try:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    except Exception:
        raise ResumeExtractionError(f"读取文本文件失败")


# ═══════════════════════════════════════
# 文本清洗
# ═══════════════════════════════════════

def _normalize_text(value: object) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.rstrip() for line in text.split("\n")).strip()


def _sanitize_resume_text(value: object, *, ext: str = "") -> str:
    """完整复刻参考项目的文本清洗逻辑"""
    text = str(value or "")
    if not text:
        return ""

    normalized = (
        text.replace("﻿", "")
        .replace("​", "")
        .replace("‌", "")
        .replace("‍", "")
        .replace("\xa0", " ")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )

    if ext in {".md", ".markdown"}:
        normalized = re.sub(r"<!--.*?-->", "", normalized, flags=re.S)
        normalized = re.sub(r"!\[[^\]]*]\([^)]+\)", "", normalized)

    lines: list[str] = []
    previous_blank = False
    for raw_line in normalized.split("\n"):
        line = raw_line.replace("\t", "    ").strip()

        if ext in {".md", ".markdown"}:
            line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)
            line = re.sub(r"^\s*>+\s*", "", line)
            line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", line)
            line = re.sub(r"`([^`]*)`", r"\1", line)
            line = line.replace("**", "").replace("__", "")

        if re.fullmatch(r"[-=_*~]{3,}", line):
            continue

        line = re.sub(r"\s+", " ", line).strip()
        if not line:
            if previous_blank:
                continue
            lines.append("")
            previous_blank = True
            continue

        lines.append(line)
        previous_blank = False

    return "\n".join(lines).strip()


def _format_direct_text_result(clean_text: str, ext: str) -> str:
    label = DIRECT_SOURCE_LABELS.get(ext, "文本文件")
    return f"【解析成功】已直接提取 {label} 内容（无需 OCR）\n\n以下是提取的内容:\n\n{clean_text}"


def unwrap_resume_text(raw_text: object, *, ext: str = "") -> str:
    return _sanitize_resume_text(_unwrap_extracted_text(str(raw_text or "")), ext=ext)


def _unwrap_extracted_text(raw_text: str) -> str:
    text = str(raw_text or "").strip()
    for marker in ("以下是提取的内容:\n\n", "以下是提取的内容：\n\n", "以下是提取的内容:", "以下是提取的内容："):
        if marker in text:
            return text.split(marker, 1)[1].strip()
    return text


def _is_successful_extraction(raw_text: str, clean_text: str) -> bool:
    if not clean_text:
        return False
    lowered_raw = raw_text.lower()
    lowered_clean = clean_text.lower()
    if any(marker in lowered_raw for marker in INVALID_TEXT_MARKERS):
        return False
    if any(marker.lower() in lowered_clean for marker in EMPTY_TEXT_MARKERS):
        return False
    return True
